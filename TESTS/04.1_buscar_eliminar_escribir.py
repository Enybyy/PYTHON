# IMPORTAR LIBRERIAS
from docx import Document
import pandas as pd

# ABRIR .DOCX
doc = Document(
    'C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/Format_contrato.docx')

print("INICIANDO...")

# INICIAR LOOP
b_e_s = ""
while b_e_s != "EXIT".lower():

    # INGRESAR DATOS A CAMBIAR
    print("========= INGRESAR DATOS A CAMBIAR =========\n")

    name = input("INGRESAR EMBAJADOR: ")
    id = input("INGRESAR DNI: ")
    address = input("INGRESAR DIRECCIÓN: ")
    district = input("INGRESAR DISTRITO: ")
    province = input("INGRESAR PROVINCIA: ")
    campaign = input("INGRESAR CAPAÑA: ")
    days = input("INGRESAR DURACIÓN: ")
    start_date = input("INGRESAR FECHA DE INICIO: ")
    end_date = input("INGRESAR FECHA DE FIN: ")
    remuneration = input("INGRESAR REMUNERACIÓN: ")

    # RECORRER CADA PARRAFO DEL DOC
    for text in doc.paragraphs:

        # CREAR LISTA text GUARDAR LOS runs FORMATEADOS
        formatted_runs = []

        # RECORRER CADA run EN EL PARRAFO
        for run in text.runs:
            # CONSERVAR EL PARRAFO OROGINAL DEL run
            formatted_run = run._element
            formatted_runs.append(formatted_run)

        if "CONTRATO" in run.text:
            print("GENERANDO .DOCX ...")
        # ELIMINAR DE LA CADENA "[NAME]" DEL PARRAFO
        text.clear()  # LIMPIAR EL CONTENIDO ORIGINAL DEL PARRAFO
        for run in formatted_runs:
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[NAME]"
            if "[NAME]" in run.text:
                text.add_run(name).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[ID]" in run.text:
                text.add_run(id).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[ADDRESS]" in run.text:
                text.add_run(address).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[DISTRICT]" in run.text:
                text.add_run(district).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[PROVINCE]" in run.text:
                text.add_run(province).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[CAMPAIGN]" in run.text:
                text.add_run(campaign).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[DAYS]" in run.text:
                text.add_run(days).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[FECHA]"
            elif "[START_DATE]" in run.text:
                text.add_run(start_date).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[END_DATE]" in run.text:
                text.add_run(end_date).bold = True
            # AGREGAR EL TEXTO PREDEFINIDO SOBRE LA POSICIÓN DE "[ID]"
            elif "[REMUNERATION]" in run.text:
                text.add_run(remuneration).bold = True
            else:
                # AGREGAR EL run ORIGINAL AL PÁRRAFO
                text._element.append(run)

    # GUARDAR DOC MODIFICADO
    try:
        doc.save(f'C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/{name}.docx')
        print(f"Documento {name}.docx generado correctamente.")
    except Exception as e:
        print(f"Error al guardar el documento: {e}")

    print(
        f"""
        El texto '[NAME]' ha reemplzado por {name}.
        El texto '[ID]' ha reemplzado por {id}.
        El texto '[ADDRESS]' ha reemplzado por {address}.
        El texto '[DISTRICT]' ha reemplzado por {district}.
        El texto '[PROVINCE]' ha reemplzado por {province}.
        El texto '[CAMPAIGN]' ha reemplzado por {campaign}.
        El texto '[DAYS]' ha reemplzado por {days}.
        El texto '[START_DATE]' ha reemplzado por {start_date}.
        El texto '[END_DATE]' ha reemplzado por {end_date}.
        El texto '[REMUNERATION]' ha reemplzado por {remuneration}.
        """)

    # TERMINAR LOOP
    b_e_s = input(
        "Si desea agregar otros datos presione 'Enter', de lo contrario escriba 'EXIT': ")
    if b_e_s == "":
        print("Saliendo del programa...")
        break
