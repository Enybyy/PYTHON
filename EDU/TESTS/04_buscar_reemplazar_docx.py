# IMPORTAR LIBRERIAS
from docx import Document
from TESTS.FUNCTIONS.swn_01 import save_words_n

# Abrir el documento de Word
doc = Document(
    'C:/Users/EVENTOS/Desktop/PROJECT_PYTHON/ARCHIVOS_LOCALES/Format_contrato.docx')

# SALVAR STRING
save_words_bl = save_words_n(doc)
save_texts = save_words_bl
# TEXT A BUSCAR Y REEMPLAZAR
find_text = "[NAME]"
replace_text = "ELIUD ROJAS MENDOZA"

# Recorrer cada párrafo del documento
for para in doc.paragraphs:

    # Verificar si la palabra a reemplazar está en el texto del párrafo
    # if find_text in para.text:
    #     # Reemplazar el texto
    #     parts = para.text.split(find_text)  # Dividimos el texto en dos partes
    #     para.clear()  # Limpiamos el contenido original del párrafo
    #     # Agregamos la primera parte del texto
    #     para.add_run(parts[0])
    #     # Agregamos la parte reemplazada en negrita
    #     para.add_run(replace_text).bold = True
    #     # Agregamos la segunda parte del texto si existe
    #     if len(parts) > 1:
    #         para.add_run(parts[1])

    # Recorrer cada palabra en el párrafo
    for run in para.runs:
        # Verificar si la palabra actual está en negrita
        if run.bold:
            # Verificar si la palabra actual está en la lista de palabras a conservar en negrita
            if any(find_text in run.text for find_text in save_texts):
                # Mantener la negrita
                continue
            else:
                # Quitar la negrita
                run.bold = False

# Guardar el documento modificado
doc.save('C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/NEW_DOCUMENT.docx')

print("Palabra reemplazada y texto en negrita en el documento de Word.")

# adsasddddddddddd-------

# Abrir el documento de Word
# doc = Document(
#     'C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/Format_contrato.docx')

# # Palabras a buscar y reemplazar
# save_words_negrita = save_words_n(doc)
# find_texts = save_words_negrita
# replace_text = "ELIUD ROJAS MENDOZA"

# # Recorrer cada párrafo del documento
# for para in doc.paragraphs:
#     # Recorrer cada palabra en el párrafo
#     for run in para.runs:
#         # Verificar si la palabra actual está en negrita
#         if run.bold:
#             # Verificar si la palabra actual está en la lista de palabras a conservar en negrita
#             if any(find_text in run.text for find_text in find_texts):
#                 # Mantener la negrita
#                 continue
#             else:
#                 # Quitar la negrita
#                 run.bold = False

# # Guardar el documento modificado
# doc.save('C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/NEW_DOCUMENT.docx')

# print("Texto en negrita conservado según especificado en el documento de Word.")


# #asddddddddddddddddddddddddd
# from docx import Document

# # Abrir el documento de Word
# doc = Document('C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/Format_contrato.docx')

# # Palabra a buscar y reemplazar
# find_text = "[NAME]"
# replace_text = "ELIUD ROJAS MENDOZA"

# # Recorrer cada párrafo del documento
# for para in doc.paragraphs:
#     # Verificar si la palabra a reemplazar está en el texto del párrafo
#     if find_text in para.text:
#         # Reemplazar el texto
#         parts = para.text.split(find_text)  # Dividimos el texto en dos partes
#         para.clear()  # Limpiamos el contenido original del párrafo
#         # Agregamos la primera parte del texto
#         para.add_run(parts[0])
#         # Agregamos la parte reemplazada en negrita
#         para.add_run(replace_text).bold = True
#         # Agregamos la segunda parte del texto si existe
#         if len(parts) > 1:
#             para.add_run(parts[1])

# # Guardar el documento modificado
# doc.save('C:/Users/EVENTOS/Desktop/PYTHON/TEST_ARCH/NEW_DOCUMENT.docx')

# print("Palabra reemplazada y texto en negrita en el documento de Word.")
